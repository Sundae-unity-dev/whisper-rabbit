"""회의 녹음 정리 docx 빌더.

JSON 한 개로 데이터를 받아 표지/목차/요약/본문 3섹션 docx 를 생성한다.
슬래시 커맨드 흐름은 회의 분류 데이터를 JSON 파일로 임시 저장하고
``python -m whisper_rabbit.report --input data.json --output out.docx``
형태로 호출하기만 하면 된다.

JSON 스키마는 :class:`MeetingData` 의 필드 그대로다. CLI ``--print-schema``
옵션으로 예시 JSON 을 stdout 에 찍어볼 수 있다.
"""
from __future__ import annotations

import argparse
import json
import logging
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Iterable, Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

log = logging.getLogger("whisper_rabbit.report")

# ─────────────────────────────────────────────────────────────
# 색상 / 상수
# ─────────────────────────────────────────────────────────────

NAVY = RGBColor(0x1F, 0x4E, 0x79)
SKY = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x60, 0x60, 0x60)
HEADER_FILL = "D9D9D9"
TLDR_FILL = "F2F2F2"
KOREAN_FONT = "맑은 고딕"

MeetingType = Literal["general", "scrum", "review", "retro"]
MetaMode = Literal["minimal", "full"]

SECTION_LABELS: dict[MeetingType, tuple[str, str, str]] = {
    "general": ("논의 내용", "문제점", "해결방안"),
    "scrum":   ("어제 한 일", "오늘 할 일", "블로커·이슈"),
    "review":  ("리뷰 의견", "개선 포인트", "후속 액션"),
    "retro":   ("Keep (잘된 것)", "Problem (개선 필요)", "Try (시도)"),
}

# 익명화 잔존 검출 키워드. 호출자가 더 추가하고 싶으면
# ReportOptions.banned_terms 로 확장 가능.
DEFAULT_BANNED_TERMS = ("슈가", "다솔림", "유지님", "지혜님", "콘치님", "제니")


# ─────────────────────────────────────────────────────────────
# 데이터 모델 (JSON 스키마)
# ─────────────────────────────────────────────────────────────

@dataclass
class KeyPoint:
    title: str
    body: str
    sub: str | None = None  # 한 줄 부연 (선택)


@dataclass
class Decision:
    title: str
    body: str


@dataclass
class ActionItem:
    number: str
    action: str
    owner: str = "?"
    due: str = "?"
    status: str = "신규"


@dataclass
class Unresolved:
    title: str
    body: str


@dataclass
class Topic:
    """본문 2/3번 섹션의 [주제] 블록. name 이 None 이면 헤더 없이 불릿 나열."""
    name: str | None
    items: list[str]


@dataclass
class MeetingData:
    """회의 분류 결과 전체. JSON 입력 스키마와 1:1 대응."""

    team: str
    date: str  # YYYY-MM-DD

    # 메타 (모두 선택. 비어 있으면 메타 표에서 생략)
    audio_file: str = ""
    audio_duration: str = ""           # "HH:MM"
    audio_sha1: str = ""
    model_label: str = ""              # "small / cpu / int8"
    transcribe_time: str = ""          # "M:SS (RTF 0.XXx)"

    # 요약 페이지
    tldr: str = ""
    keywords: list[str] = field(default_factory=list)
    key_points: list[KeyPoint] = field(default_factory=list)
    decisions: list[Decision] = field(default_factory=list)
    actions: list[ActionItem] = field(default_factory=list)
    unresolved: list[Unresolved] = field(default_factory=list)
    context: list[str] = field(default_factory=list)

    # 본문 3 섹션. body[0] = 섹션 1, body[1] = 2, body[2] = 3
    # 각 섹션은 Topic 의 list. Topic.name=None 이면 평이 불릿.
    body: list[list[Topic]] = field(default_factory=lambda: [[], [], []])

    # ─ 직렬화 ─
    @classmethod
    def from_dict(cls, d: dict[str, Any]) -> "MeetingData":
        return cls(
            team=d["team"],
            date=d["date"],
            audio_file=d.get("audio_file", ""),
            audio_duration=d.get("audio_duration", ""),
            audio_sha1=d.get("audio_sha1", ""),
            model_label=d.get("model_label", ""),
            transcribe_time=d.get("transcribe_time", ""),
            tldr=d.get("tldr", ""),
            keywords=list(d.get("keywords", [])),
            key_points=[KeyPoint(**x) for x in d.get("key_points", [])],
            decisions=[Decision(**x) for x in d.get("decisions", [])],
            actions=[ActionItem(**x) for x in d.get("actions", [])],
            unresolved=[Unresolved(**x) for x in d.get("unresolved", [])],
            context=list(d.get("context", [])),
            body=[
                [Topic(**t) for t in section] for section in d.get("body", [[], [], []])
            ],
        )

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class ReportOptions:
    meeting_type: MeetingType = "general"
    meta_mode: MetaMode = "minimal"
    include_cover: bool = True
    include_toc: bool = True
    include_context: bool = True
    banned_terms: tuple[str, ...] = DEFAULT_BANNED_TERMS

    @property
    def section_labels(self) -> tuple[str, str, str]:
        return SECTION_LABELS[self.meeting_type]


# ─────────────────────────────────────────────────────────────
# 스타일 헬퍼 (run / cell)
# ─────────────────────────────────────────────────────────────

def set_korean_font(
    run,
    size_pt: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = KOREAN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(k), KOREAN_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def style_cell_text(
    cell,
    text: str,
    size_pt: float = 10.5,
    bold: bool = False,
    color: RGBColor | None = None,
    align: int | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_korean_font(run, size_pt=size_pt, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ─────────────────────────────────────────────────────────────
# 단락 / 표 / 필드 헬퍼
# ─────────────────────────────────────────────────────────────

def add_bullet(
    doc: Document,
    text: str,
    bold_prefix: str | None = None,
    indent_inches: float = 0.0,
    size_pt: float = 11.0,
) -> None:
    style = "List Bullet" if indent_inches == 0 else "Normal"
    p = doc.add_paragraph(style=style)
    if indent_inches:
        p.paragraph_format.left_indent = Inches(indent_inches)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_korean_font(r1, size_pt=size_pt, bold=True)
        if text:
            r2 = p.add_run(" — " + text)
            set_korean_font(r2, size_pt=size_pt)
    else:
        prefix = "· " if indent_inches else ""
        r = p.add_run(prefix + text)
        set_korean_font(r, size_pt=size_pt)


def add_numbered(doc: Document, idx: int, bold_prefix: str, body: str) -> None:
    p = doc.add_paragraph()
    r0 = p.add_run(f"{idx}. ")
    set_korean_font(r0, size_pt=11, bold=True)
    r1 = p.add_run(bold_prefix)
    set_korean_font(r1, size_pt=11, bold=True)
    r2 = p.add_run(" — " + body)
    set_korean_font(r2, size_pt=11)


def add_section_heading(doc: Document, text: str, force_page_break: bool = True) -> None:
    """Heading 1 + 한국어 폰트·색상 + (옵션) 페이지 직전 강제 줄바꿈."""
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    if force_page_break:
        pPr = p._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:pageBreakBefore"))
    r = p.add_run(text)
    set_korean_font(r, size_pt=15, bold=True, color=NAVY)


def add_sub_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"[{text}]")
    set_korean_font(r, size_pt=12, bold=True, color=SKY)


def add_tldr_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E75B6")
    pbdr.append(left)
    pPr.append(pbdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), TLDR_FILL)
    pPr.append(shd)
    r1 = p.add_run("TL;DR  ")
    set_korean_font(r1, size_pt=11, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_korean_font(r2, size_pt=11)


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    style_cell_text(table.cell(0, 0), "항목", size_pt=10.5, bold=True)
    style_cell_text(table.cell(0, 1), "값", size_pt=10.5, bold=True)
    shade_cell(table.cell(0, 0), HEADER_FILL)
    shade_cell(table.cell(0, 1), HEADER_FILL)
    for i, (k, v) in enumerate(rows, start=1):
        style_cell_text(table.cell(i, 0), k, size_pt=10.5, bold=True)
        style_cell_text(table.cell(i, 1), v, size_pt=10.5)


def add_keywords_line(doc: Document, kws: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r0 = p.add_run("핵심 키워드  ")
    set_korean_font(r0, size_pt=11, bold=True, color=NAVY)
    r1 = p.add_run(" · ".join(kws))
    set_korean_font(r1, size_pt=11)


def add_action_table(doc: Document, rows: Iterable[ActionItem]) -> None:
    rows = list(rows)
    headers = ["#", "액션", "담당", "기한", "상태/비고"]
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    widths = [Cm(1.0), Cm(9.5), Cm(1.5), Cm(1.5), Cm(2.5)]
    for col_idx, w in enumerate(widths):
        for r in table.rows:
            r.cells[col_idx].width = w
    for i, h in enumerate(headers):
        align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else None
        style_cell_text(table.cell(0, i), h, size_pt=10.5, bold=True, align=align)
        shade_cell(table.cell(0, i), HEADER_FILL)
    for ridx, item in enumerate(rows, start=1):
        cells_vals = (item.number, item.action, item.owner, item.due, item.status)
        for cidx, val in enumerate(cells_vals):
            align = WD_ALIGN_PARAGRAPH.CENTER if cidx in (0, 2, 3, 4) else None
            style_cell_text(table.cell(ridx, cidx), val, size_pt=10.5, align=align)


def add_toc_field(doc: Document) -> None:
    """Heading 1·2 자동 수집 TOC. updateFields=true 와 함께 쓰면 Word 열 때 자동 채움."""
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    pt = OxmlElement("w:t")
    pt.text = "(목차 자동 생성 중…)"
    placeholder.append(pt)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._element.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    set_korean_font(r, size_pt=9, color=GREY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._element.extend([fld_begin, instr, fld_sep, placeholder_t, fld_end])


def enable_auto_update_fields(doc: Document) -> None:
    """Word 가 처음 열 때 모든 필드(TOC/PAGE) 자동 업데이트하도록 settings.xml 에 신호."""
    settings = doc.settings.element
    for child in settings:
        if child.tag == qn("w:updateFields"):
            child.set(qn("w:val"), "true")
            return
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ─────────────────────────────────────────────────────────────
# 검증
# ─────────────────────────────────────────────────────────────

def verify_anonymization(doc: Document, banned: tuple[str, ...]) -> list[str]:
    """문서 전체에서 금지 키워드(실명/단정 표현) 잔존 여부 검사."""
    violations: list[str] = []
    for p in doc.paragraphs:
        for kw in banned:
            if kw in p.text:
                violations.append(f"paragraph: '{kw}' in '{p.text[:80]}'")
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for kw in banned:
                    if kw in cell.text:
                        violations.append(f"table cell: '{kw}' in '{cell.text[:80]}'")
    return violations


# ─────────────────────────────────────────────────────────────
# 본 빌더
# ─────────────────────────────────────────────────────────────

def _set_normal_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(k), KOREAN_FONT)
    rPr.insert(0, rFonts)


def _meta_rows(data: MeetingData, options: ReportOptions, section_counts: tuple[int, int, int]) -> list[tuple[str, str]]:
    s1, s2, s3 = options.section_labels
    n1, n2, n3 = section_counts
    classification = f"{s1} {n1} / {s2} {n2} / {s3} {n3}"
    action_summary = (
        f"{len(data.actions)}건 "
        f"(담당 미지정 {sum(1 for a in data.actions if a.owner == '?')} "
        f"/ 기한 미지정 {sum(1 for a in data.actions if a.due == '?')})"
    )
    minimal = [
        ("녹음 길이", data.audio_duration or "—"),
        ("분류 항목", classification),
        ("액션 아이템", action_summary),
    ]
    if options.meta_mode == "minimal":
        return minimal
    extra = [
        ("받아쓰기", data.transcribe_time or "—"),
        ("모델 / 장치", data.model_label or "—"),
        ("audio_sha1", data.audio_sha1 or "—"),
    ]
    return minimal + extra


def build_docx(data: MeetingData, options: ReportOptions, out_path: Path) -> Path:
    """회의록 docx 를 만들고 저장 경로를 반환.

    PermissionError(파일이 Word 등에서 열려 있음) 발생 시 ``_v2``, ``_v3`` 로
    증분 저장.
    """
    doc = Document()
    _set_normal_font(doc)
    enable_auto_update_fields(doc)
    add_page_number_footer(doc)

    s1, s2, s3 = options.section_labels
    section_counts = (
        sum(len(t.items) for t in data.body[0]) if len(data.body) > 0 else 0,
        sum(len(t.items) for t in data.body[1]) if len(data.body) > 1 else 0,
        sum(len(t.items) for t in data.body[2]) if len(data.body) > 2 else 0,
    )

    # 표지
    if options.include_cover:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(80)
        r = title_p.add_run(f"{data.team} 회의 녹음 정리")
        set_korean_font(r, size_pt=20, bold=True)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub_p.add_run(f"({data.date})")
        set_korean_font(r, size_pt=14)

        cover_lines = []
        if data.audio_file:
            cover_lines.append(f"원본 오디오: {data.audio_file}")
        if data.audio_duration:
            cover_lines.append(f"녹음 길이: {data.audio_duration}")
        cover_lines.append(f"회의 유형: {options.meeting_type}")
        if data.transcribe_time:
            cover_lines.append(f"받아쓰기 처리: {data.transcribe_time}")
        for line in cover_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            set_korean_font(r, size_pt=10, color=GREY)

    # 목차
    if options.include_toc:
        title_p = doc.add_paragraph()
        if options.include_cover:
            pPr = title_p._p.get_or_add_pPr()
            pPr.append(OxmlElement("w:pageBreakBefore"))
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_r = title_p.add_run("목차")
        set_korean_font(title_r, size_pt=18, bold=True, color=NAVY)
        add_toc_field(doc)

    # 0. 회의 요약
    add_section_heading(
        doc, "■ 0. 회의 요약",
        force_page_break=options.include_cover or options.include_toc,
    )
    if data.tldr:
        add_tldr_block(doc, data.tldr)
    doc.add_paragraph()
    add_meta_table(doc, _meta_rows(data, options, section_counts))
    if data.keywords:
        add_keywords_line(doc, data.keywords)

    if data.key_points:
        add_sub_heading(doc, "핵심 포인트")
        for kp in data.key_points:
            add_bullet(doc, kp.body, bold_prefix=kp.title)
            if kp.sub:
                add_bullet(doc, kp.sub, indent_inches=0.4, size_pt=10)

    if data.decisions:
        add_sub_heading(doc, "결정 사항")
        for i, dec in enumerate(data.decisions, start=1):
            add_numbered(doc, i, dec.title, dec.body)

    if data.actions:
        add_sub_heading(doc, "액션 아이템")
        add_action_table(doc, data.actions)
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        r = note.add_run(
            "담당·기한이 transcript 에서 확정되지 않아 ‘?’로 둠. "
            "회의 직후 수기로 채워 넣는 것을 권장."
        )
        set_korean_font(r, size_pt=9, color=GREY)

    if data.unresolved:
        add_sub_heading(doc, "미해결 이슈")
        for u in data.unresolved:
            add_bullet(doc, u.body, bold_prefix=u.title)

    if options.include_context and data.context:
        add_sub_heading(doc, "회의 맥락")
        for item in data.context:
            add_bullet(doc, item)

    # 본문 3 섹션
    body_titles = (f"■ 1. {s1}", f"■ 2. {s2}", f"■ 3. {s3}")
    for idx, title in enumerate(body_titles):
        topics = data.body[idx] if idx < len(data.body) else []
        add_section_heading(doc, title, force_page_break=True)
        for topic in topics:
            if topic.name:
                add_sub_heading(doc, topic.name)
            for line in topic.items:
                add_bullet(doc, line)

    # 검증
    violations = verify_anonymization(doc, options.banned_terms)
    if violations:
        for v in violations:
            log.error("[VIOLATION] %s", v)
        raise RuntimeError(
            f"익명화 검증 실패: {len(violations)}건 잔존. 입력 JSON 수정 필요."
        )

    # 증분 저장
    candidate = out_path
    i = 2
    while True:
        try:
            doc.save(candidate)
            return candidate
        except PermissionError:
            candidate = candidate.with_name(
                f"{out_path.stem}_v{i}{out_path.suffix}"
            )
            i += 1
            if i > 10:
                raise


# ─────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────

_EXAMPLE_JSON = {
    "team": "팀명",
    "date": "2026-05-13",
    "audio_file": "meeting.mp4",
    "audio_duration": "13:23",
    "audio_sha1": "abc123...",
    "model_label": "small / cpu / int8",
    "transcribe_time": "3:44 (RTF 0.28x)",
    "tldr": "회의 결론 1~3문장.",
    "keywords": ["키워드1", "키워드2"],
    "key_points": [
        {"title": "포인트 제목", "body": "한 줄 본문", "sub": "선택 부연"},
    ],
    "decisions": [
        {"title": "결정 제목", "body": "내용 + 이유"},
    ],
    "actions": [
        {"number": "1", "action": "할 일", "owner": "?", "due": "?", "status": "신규"},
    ],
    "unresolved": [
        {"title": "이슈 제목", "body": "사유 + 다음 결정 시점"},
    ],
    "context": ["출결·휴가 등 의제 외 정보"],
    "body": [
        [{"name": None, "items": ["섹션 1 불릿 1", "섹션 1 불릿 2"]}],
        [{"name": "주제 키워드", "items": ["섹션 2 불릿 1"]}],
        [{"name": "주제 키워드", "items": ["섹션 3 불릿 1"]}],
    ],
}


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="whisper_rabbit.report",
        description="JSON 으로 받은 회의 분류 결과를 회의록 docx 로 변환.",
    )
    p.add_argument("--input", type=Path, help="회의 데이터 JSON 경로")
    p.add_argument("--output", type=Path,
                   help="출력 docx 경로. 미지정 시 Desktop/<팀>_회의녹음정리_<날짜>.docx")
    p.add_argument("--meeting-type", default="general",
                   choices=["general", "scrum", "review", "retro"])
    p.add_argument("--meta", default="minimal", choices=["minimal", "full"])
    p.add_argument("--no-cover", action="store_true")
    p.add_argument("--no-toc", action="store_true")
    p.add_argument("--no-context", action="store_true")
    p.add_argument("--print-schema", action="store_true",
                   help="예시 JSON 을 stdout 에 출력하고 종료")
    p.add_argument("-v", "--verbose", action="count", default=0)
    return p


def _default_output(data: MeetingData) -> Path:
    desktop = Path.home() / "Desktop"
    desktop.mkdir(parents=True, exist_ok=True)
    return desktop / f"{data.team}_회의녹음정리_{data.date}.docx"


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    logging.basicConfig(
        level=logging.DEBUG if args.verbose >= 2 else logging.INFO,
        format="%(asctime)s [%(levelname).1s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
    )

    if args.print_schema:
        print(json.dumps(_EXAMPLE_JSON, ensure_ascii=False, indent=2))
        return 0

    if not args.input:
        print("ERROR: --input <json> 필수 (또는 --print-schema)", file=sys.stderr)
        return 2
    if not args.input.exists():
        print(f"ERROR: 입력 JSON 없음: {args.input}", file=sys.stderr)
        return 2

    raw = json.loads(args.input.read_text(encoding="utf-8"))
    data = MeetingData.from_dict(raw)
    options = ReportOptions(
        meeting_type=args.meeting_type,
        meta_mode=args.meta,
        include_cover=not args.no_cover,
        include_toc=not args.no_toc,
        include_context=not args.no_context,
    )
    out_path = args.output or _default_output(data)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    saved = build_docx(data, options, out_path)
    summary = {
        "docx": str(saved),
        "docx_size_kb": round(saved.stat().st_size / 1024, 1),
        "team": data.team,
        "date": data.date,
        "options": {
            "meeting_type": options.meeting_type,
            "meta_mode": options.meta_mode,
            "include_cover": options.include_cover,
            "include_toc": options.include_toc,
            "include_context": options.include_context,
        },
        "section_counts": {
            "key_points": len(data.key_points),
            "decisions": len(data.decisions),
            "actions": len(data.actions),
            "unresolved": len(data.unresolved),
            "context": len(data.context),
            options.section_labels[0]: sum(len(t.items) for t in (data.body[0] if len(data.body) > 0 else [])),
            options.section_labels[1]: sum(len(t.items) for t in (data.body[1] if len(data.body) > 1 else [])),
            options.section_labels[2]: sum(len(t.items) for t in (data.body[2] if len(data.body) > 2 else [])),
        },
        "verification": "passed",
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
