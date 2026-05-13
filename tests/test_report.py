from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from whisper_rabbit.report import (
    ActionItem,
    Decision,
    KeyPoint,
    MeetingData,
    ReportOptions,
    Topic,
    Unresolved,
    _meta_rows,
    build_docx,
    verify_anonymization,
)


@pytest.fixture
def sample_data() -> MeetingData:
    return MeetingData(
        team="팀A",
        date="2026-05-13",
        audio_file="x.mp4",
        audio_duration="13:23",
        audio_sha1="deadbeef",
        model_label="small / cpu / int8",
        transcribe_time="3:44 (RTF 0.28x)",
        tldr="짧은 회의 요약",
        keywords=["k1", "k2"],
        key_points=[
            KeyPoint(title="p1", body="b1", sub="s1"),
            KeyPoint(title="p2", body="b2"),
        ],
        decisions=[Decision(title="d1", body="이유1")],
        actions=[
            ActionItem(number="1", action="할 일 1", owner="?", due="?", status="신규"),
            ActionItem(number="2", action="할 일 2", owner="홍길동", due="2026-05-20", status="진행 중"),
        ],
        unresolved=[Unresolved(title="이슈1", body="사유")],
        context=["출결 항목 1"],
        body=[
            [Topic(name=None, items=["불릿 1", "불릿 2"])],
            [Topic(name="주제A", items=["문제 1"])],
            [Topic(name="주제A", items=["해결 1"])],
        ],
    )


class TestSerialization:
    def test_from_dict_to_dict_roundtrip(self, sample_data: MeetingData) -> None:
        d = sample_data.to_dict()
        # JSON 직렬화 가능 (nested dataclass 가 dict 로 풀려야 함)
        s = json.dumps(d, ensure_ascii=False)
        roundtrip = MeetingData.from_dict(json.loads(s))
        assert roundtrip.team == sample_data.team
        assert len(roundtrip.key_points) == 2
        assert roundtrip.key_points[0].title == "p1"
        assert roundtrip.actions[1].owner == "홍길동"
        assert roundtrip.body[1][0].name == "주제A"

    def test_minimal_payload(self) -> None:
        d = MeetingData.from_dict({"team": "T", "date": "2026-05-13"})
        assert d.team == "T"
        assert d.tldr == ""
        assert d.key_points == []
        assert d.body == [[], [], []]


class TestMetaRows:
    def test_minimal_has_three_rows(self, sample_data: MeetingData) -> None:
        opt = ReportOptions(meta_mode="minimal")
        rows = _meta_rows(sample_data, opt, (2, 1, 1))
        assert len(rows) == 3
        keys = [k for k, _ in rows]
        assert keys == ["녹음 길이", "분류 항목", "액션 아이템"]

    def test_full_adds_three_more(self, sample_data: MeetingData) -> None:
        opt = ReportOptions(meta_mode="full")
        rows = _meta_rows(sample_data, opt, (2, 1, 1))
        assert len(rows) == 6
        assert "audio_sha1" in [k for k, _ in rows]

    def test_classification_uses_section_labels(self, sample_data: MeetingData) -> None:
        opt = ReportOptions(meeting_type="scrum", meta_mode="minimal")
        rows = _meta_rows(sample_data, opt, (3, 4, 5))
        classification = dict(rows)["분류 항목"]
        assert "어제 한 일 3" in classification
        assert "오늘 할 일 4" in classification
        assert "블로커·이슈 5" in classification

    def test_action_summary_counts_unassigned(self, sample_data: MeetingData) -> None:
        opt = ReportOptions(meta_mode="minimal")
        rows = _meta_rows(sample_data, opt, (1, 1, 1))
        action = dict(rows)["액션 아이템"]
        # 샘플: 2건, 담당 미지정 1, 기한 미지정 1
        assert "2건" in action
        assert "담당 미지정 1" in action
        assert "기한 미지정 1" in action


class TestBuild:
    def test_build_produces_docx(self, tmp_path: Path, sample_data: MeetingData) -> None:
        out = tmp_path / "out.docx"
        saved = build_docx(sample_data, ReportOptions(), out)
        assert saved == out
        assert saved.exists()
        assert saved.stat().st_size > 5_000  # 의미 있는 크기

    def test_meeting_type_changes_section_labels(self, tmp_path: Path, sample_data: MeetingData) -> None:
        for mt, expected_labels in [
            ("general", ("논의 내용", "문제점", "해결방안")),
            ("scrum", ("어제 한 일", "오늘 할 일", "블로커·이슈")),
            ("retro", ("Keep (잘된 것)", "Problem (개선 필요)", "Try (시도)")),
        ]:
            out = tmp_path / f"{mt}.docx"
            build_docx(sample_data, ReportOptions(meeting_type=mt), out)  # type: ignore[arg-type]
            doc = Document(str(out))
            text = "\n".join(p.text for p in doc.paragraphs)
            for label in expected_labels:
                assert label in text, f"{mt}: 라벨 '{label}' 누락"

    def test_no_cover_no_toc_no_context_skips(self, tmp_path: Path, sample_data: MeetingData) -> None:
        out = tmp_path / "minimal.docx"
        build_docx(
            sample_data,
            ReportOptions(include_cover=False, include_toc=False, include_context=False),
            out,
        )
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "회의 맥락" not in text
        assert "목차" not in text
        # 표지 부제는 "(YYYY-MM-DD)" 형식 — 없어야 함
        assert "(2026-05-13)" not in text


class TestVerifyAnonymization:
    def test_passes_clean(self, tmp_path: Path, sample_data: MeetingData) -> None:
        out = tmp_path / "ok.docx"
        # 검증은 build_docx 내부에서 실행 — 통과해야 saved 됨
        build_docx(sample_data, ReportOptions(), out)
        # 추가로 직접 호출도 통과
        violations = verify_anonymization(Document(str(out)), ("슈가", "다솔림"))
        assert violations == []

    def test_detects_banned_word(self, tmp_path: Path) -> None:
        bad = MeetingData(
            team="T", date="2026-05-13",
            key_points=[KeyPoint(title="t", body="슈가 사용 정리")],
        )
        out = tmp_path / "bad.docx"
        with pytest.raises(RuntimeError, match="익명화"):
            build_docx(bad, ReportOptions(banned_terms=("슈가",)), out)
